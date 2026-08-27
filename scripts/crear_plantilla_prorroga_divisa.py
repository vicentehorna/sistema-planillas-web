# -*- coding: utf-8 -*-
"""
Convierte PRORROGA 1_VIAMERICA.doc a plantilla .docx web
con marcadores Jinja (docxtpl) para hm_divisa (cia BGT).
Preserva la firma/imagen del empleador del documento original.
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

import pythoncom
import win32com.client  # type: ignore
from docx import Document

ROOT = Path(r"C:\PROYECTOS\PLANILLAS")
sys.path.insert(0, str(ROOT))
from contratos_docs import MARCADORES_LEGACY_MAP, plantillas_dir  # noqa: E402

SRC = ROOT / "AUXILIARES" / "CONTRATOS" / "PRORROGA 1_VIAMERICA.doc"
OUT_AUX = ROOT / "AUXILIARES" / "CONTRATOS" / "PRORROGA_1_VIAMERICA.docx"
CIA = "BGT"

PRORROGA_LEGACY = {
    "NOMBRETRAB": "trabajador",
    "NRODOCTRAB": "dni",
    "DIRECCIONTRAB1": "direccion",
    "FECHAINIANTERIOR": "inicio_contrato_anterior",
    "CARGOTRAB": "cargo",
    "FECHAFINANTERIOR": "fin_contrato_anterior",
    "PERIODOCONTRATO": "periodo_contrato",
    "FECHAINICONTRATO": "inicio_contrato",
    "FECHAFINCONTRATO": "fin_contrato",
    "FIRMACONTRATOTEXTO1": "fecha_firma",
}


def _wd_format_docx():
    return 12


def convert_and_replace_bookmarks() -> Path:
    pythoncom.CoInitialize()
    word = None
    doc = None
    tmp_docx = ROOT / "AUXILIARES" / "CONTRATOS" / "_tmp_prorroga_divisa.docx"
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(
            FileName=str(SRC),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
        )
        doc.SaveAs2(FileName=str(tmp_docx), FileFormat=_wd_format_docx())
        doc.Close(False)
        doc = None
        doc = word.Documents.Open(
            FileName=str(tmp_docx),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
        )

        mapping = {**MARCADORES_LEGACY_MAP, **PRORROGA_LEGACY}
        names = [b.Name for b in doc.Bookmarks]
        print("Bookmarks encontrados:", names)
        for bk_name in list(names):
            if str(bk_name).startswith("_"):
                continue
            web = mapping.get(str(bk_name).upper())
            if not web:
                print(f"  SKIP sin mapa: {bk_name}")
                continue
            marker = "{{ " + web + " }}"
            try:
                if doc.Bookmarks.Exists(bk_name):
                    doc.Bookmarks(bk_name).Range.Text = marker
                    print(f"  OK {bk_name} -> {marker}")
            except Exception as ex:
                print(f"  ERR {bk_name}: {ex}")

        print(f"Shapes (firma): {doc.Shapes.Count}")

        doc.SaveAs2(FileName=str(OUT_AUX), FileFormat=_wd_format_docx())
        doc.Close(False)
        doc = None
        word.Quit()
        word = None
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()
        if tmp_docx.exists():
            try:
                tmp_docx.unlink()
            except Exception:
                pass
    return OUT_AUX


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        for r in paragraph.runs:
            r.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _collect_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def enrich_with_web_markers(path: Path) -> None:
    doc = Document(str(path))
    joined = _collect_text(doc)
    present = set(re.findall(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}", joined))
    print("Marcadores en plantilla:", sorted(present))

    legacy_names = list({*MARCADORES_LEGACY_MAP.keys(), *PRORROGA_LEGACY.keys()})
    for p in doc.paragraphs:
        t = p.text
        orig = t
        for leg in legacy_names:
            t = re.sub(rf"\b{re.escape(leg)}\b\s*", "", t)
        t = re.sub(r"[ \t]{2,}", " ", t)
        if t != orig:
            _set_paragraph_text(p, t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    orig = t
                    for leg in legacy_names:
                        t = re.sub(rf"\b{re.escape(leg)}\b\s*", "", t)
                    t = re.sub(r"[ \t]{2,}", " ", t)
                    if t != orig:
                        _set_paragraph_text(p, t)

    doc.save(str(path))
    joined2 = _collect_text(Document(str(path)))
    present2 = sorted(set(re.findall(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}", joined2)))
    print("Marcadores finales:", present2)

    expected = sorted(set(PRORROGA_LEGACY.values()))
    missing = [m for m in expected if m not in present2]
    if missing:
        raise SystemExit(f"Faltan marcadores en plantilla: {missing}")


def install_for_cia(path: Path, cia: str) -> Path:
    dest_dir = plantillas_dir(cia)
    dest = dest_dir / "PRORROGA_1_VIAMERICA.docx"
    shutil.copy2(path, dest)
    print("Instalada en", dest)
    return dest


def main():
    if not SRC.is_file():
        raise SystemExit(f"No existe: {SRC}")
    out = convert_and_replace_bookmarks()
    enrich_with_web_markers(out)
    install_for_cia(out, CIA)
    print("LISTO:", out)


if __name__ == "__main__":
    main()
